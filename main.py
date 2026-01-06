import tkinter as tk
from tkinter import filedialog, messagebox, ttk, simpledialog
import pandas as pd
import random
import os
import json
import configparser
import sys
# 添加PIL库导入
try:
    from PIL import Image, ImageDraw, ImageFont
except ImportError:
    print("未找到PIL库，请先安装: pip install pillow")
    pillow_available = False
else:
    pillow_available = True

# 尝试导入需要的库
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("未找到python-docx库，请先安装: pip install python-docx")
    docx_available = False
else:
    docx_available = True


class StudentSeatTool:
    def __init__(self, root):
        self.root = root
        
        # 读取配置文件
        self.config = self.load_config()
        
        # 设置应用基本参数（硬编码默认值）
        self.root.title("学生座位表调整工具 v1.0 By:侯小圣")
        self.root.geometry("1240x800")
        self.root.configure(bg="#f8f8f8")  # 更柔和的背景色
        # 添加窗口图标
        try:
            img_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "images", "img.ico")
            if os.path.exists(img_path):
                self.root.iconbitmap(img_path)
        except:
            pass  # 忽略图标加载错误
        # 固定窗口大小，不允许用户调整
        self.root.resizable(False, False)

        # 布局配置（从配置文件读取值）
        self.layout_config = {
            "podium_seats": self.config.getint("Layout", "podium_seats", fallback=0),  # 讲台侧座位数
            "main_rows": self.config.getint("Layout", "main_rows", fallback=6),  # 主体座位行数
            "main_cols": self.config.getint("Layout", "main_cols", fallback=8),  # 主体座位列数
            "class_name": self.config.get("Layout", "class_name", fallback=""),  # 班级名称
            "teacher_name": self.config.get("Layout", "teacher_name", fallback="")  # 班主任姓名
        }
        self.seat_positions = []  # 动态生成的座位坐标
        self.seat_index_map = {}  # 座位编号映射
        # 性别颜色（改进的配色方案，更加柔和美观）
        self.gender_color = {
            "男": "#64B5F6",  # 柔和的蓝色
            "女": "#FFB7C5",  # 柔和的粉色
            "空": "#E8EAF6"   # 柔和的灰色
        }
        self.students = []
        self.seat_data = {}
        self.drag_source = None
        self.seat_buttons = {}
        
        # 创建ttk样式，用于实现圆角效果
        self.style = ttk.Style()
        # 配置座位框架样式
        self.style.configure("RoundedFrame.TLabelframe", 
                             borderwidth=2, 
                             relief="solid",
                             padding=15)
        self.style.configure("RoundedFrame.TLabelframe.Label", 
                             font=("微软雅黑", 12, "bold"))
        # 设置Labelframe的背景色
        self.style.configure("RoundedFrame.TLabelframe", 
                             background="#f8f8f8")

        # 构建UI
        self.create_header()
        self.create_toolbar()
        self.create_seat_container()
        # 初始化座位布局
        self.generate_seat_positions()
        # 程序启动时自动尝试读取数据文件
        self.auto_load_data()

    def load_config(self):
        """加载配置文件，仅读取班级信息和座位行列配置"""
        config = configparser.ConfigParser()
        
        # 获取程序运行时的基础目录（考虑PyInstaller单文件打包情况）
        if hasattr(sys, '_MEIPASS'):
            # 打包后的临时目录
            base_dir = os.path.dirname(os.path.abspath(sys.executable))
        else:
            # 正常运行时的目录
            base_dir = os.path.dirname(os.path.abspath(__file__))
            
        config_file = os.path.join(base_dir, "config.ini")
        
        # 如果配置文件存在，读取配置
        if os.path.exists(config_file):
            config.read(config_file, encoding='utf-8')
        else:
            # 如果配置文件不存在，创建默认配置
            self.create_default_config(config, config_file)
        
        return config
        
    def parse_tuple_str(self, tuple_str):
        """安全地将元组字符串转换为元组
        
        Args:
            tuple_str: 形如'(x,y)'的字符串
            
        Returns:
            tuple: 解析后的(x,y)元组，解析失败返回None
        """
        try:
            # 去除元组的括号
            inner = tuple_str.strip('()')
            # 分割数字部分
            parts = inner.split(',')
            # 转换为整数并返回元组
            return (int(parts[0].strip()), int(parts[1].strip()))
        except (ValueError, IndexError) as e:
            print(f"解析元组时出错: {e}")
            return None
    
    def create_default_config(self, config, config_file):
        """创建默认配置文件，仅包含班级信息和座位行列配置"""
        config['Layout'] = {
            'class_name': '',
            'teacher_name': '',
            'podium_seats': '0',
            'main_rows': '6',
            'main_cols': '8'
        }
        
        # 保存配置文件
        with open(config_file, 'w', encoding='utf-8') as f:
            config.write(f)
    
    def create_header(self):
        # 改进的标题栏设计
        header_frame = tk.Frame(self.root, bg="#2196F3", height=65, bd=0, relief=tk.FLAT)
        header_frame.pack(fill=tk.X, anchor="n")
        
        # 添加标题左侧装饰条
        decor_frame = tk.Frame(header_frame, bg="#1976D2", width=5, height=65)
        decor_frame.pack(side=tk.LEFT, fill=tk.Y)
        
        # 标题标签
        tk.Label(
            header_frame, text="学生座位表调整工具",
            font=("微软雅黑", 21, "bold"), fg="white", bg="#2196F3"
        ).pack(pady=10, side=tk.LEFT, padx=15)

    def create_toolbar(self):
        toolbar_frame = tk.Frame(self.root, bg="#f5f5f5", padx=8, pady=4)
        toolbar_frame.pack(fill=tk.X, anchor="n")

        # 创建左侧按钮组
        left_frame = tk.Frame(toolbar_frame, bg="#f5f5f5")
        left_frame.pack(side=tk.LEFT, fill=tk.Y)

        # 导入按钮
        self.import_btn = tk.Button(
            left_frame, text="导入Excel", bg="#2196F3", fg="white",
            font=("微软雅黑", 11), padx=10, pady=4, relief=tk.RAISED, bd=2, command=self.import_excel
        )
        self.import_btn.pack(side=tk.LEFT, padx=3)

        # 分隔线
        tk.Frame(left_frame, width=2, bg="#ddd").pack(side=tk.LEFT, fill=tk.Y, padx=5, pady=2)

        # 排列按钮组
        arrange_frame = tk.Frame(left_frame, bg="#f5f5f5")
        arrange_frame.pack(side=tk.LEFT, fill=tk.Y)

        self.random_btn = tk.Button(
            arrange_frame, text="随机排列", bg="#FFC107", fg="black",
            font=("微软雅黑", 11), padx=10, pady=4, relief=tk.RAISED, bd=2, command=self.random_arrange
        )
        self.random_btn.pack(side=tk.LEFT, padx=3)

        self.height_btn = tk.Button(
            arrange_frame, text="按身高排序", bg="#4CAF50", fg="white",
            font=("微软雅黑", 11), padx=10, pady=4, relief=tk.RAISED, bd=2, command=self.sort_by_height
        )
        self.height_btn.pack(side=tk.LEFT, padx=3)
        
        self.score_btn = tk.Button(
            arrange_frame, text="按成绩排序", bg="#FF9800", fg="white",
            font=("微软雅黑", 11), padx=10, pady=4, relief=tk.RAISED, bd=2, command=self.sort_by_score
        )
        self.score_btn.pack(side=tk.LEFT, padx=3)

        # 分隔线
        tk.Frame(left_frame, width=2, bg="#ddd").pack(side=tk.LEFT, fill=tk.Y, padx=5, pady=2)

        # 操作按钮组
        action_frame = tk.Frame(left_frame, bg="#f5f5f5")
        action_frame.pack(side=tk.LEFT, fill=tk.Y)

        self.reset_btn = tk.Button(
            action_frame, text="重置座位", bg="#F44336", fg="white",
            font=("微软雅黑", 11), padx=10, pady=4, relief=tk.RAISED, bd=2, command=self.reset_seats
        )
        self.reset_btn.pack(side=tk.LEFT, padx=3)
        
        self.save_btn = tk.Button(
            action_frame, text="保存数据", bg="#4CAF50", fg="white",
            font=("微软雅黑", 11), padx=10, pady=4, relief=tk.RAISED, bd=2, command=self.save_data
        )
        self.save_btn.pack(side=tk.LEFT, padx=3)
        
        self.load_btn = tk.Button(
            action_frame, text="加载数据", bg="#FF9800", fg="white",
            font=("微软雅黑", 11), padx=10, pady=4, relief=tk.RAISED, bd=2, command=self.load_data
        )
        self.load_btn.pack(side=tk.LEFT, padx=3)

        # 分隔线
        tk.Frame(left_frame, width=2, bg="#ddd").pack(side=tk.LEFT, fill=tk.Y, padx=5, pady=2)

        # 导出按钮组
        export_frame = tk.Frame(left_frame, bg="#f5f5f5")
        export_frame.pack(side=tk.LEFT, fill=tk.Y)

        self.export_pdf_btn = tk.Button(
            export_frame, text="导出PDF", bg="#F44336", fg="white",
            font=("微软雅黑", 11), padx=10, pady=4, relief=tk.RAISED, bd=2, command=self.export_pdf
        )
        self.export_pdf_btn.pack(side=tk.LEFT, padx=3)
        
        self.export_docx_btn = tk.Button(
            export_frame, text="导出Word", bg="#9C27B0", fg="white",
            font=("微软雅黑", 11), padx=10, pady=4, relief=tk.RAISED, bd=2, command=self.export_layout_to_word
        )
        self.export_docx_btn.pack(side=tk.LEFT, padx=3)

        # 右侧按钮
        right_frame = tk.Frame(toolbar_frame, bg="#f5f5f5")
        right_frame.pack(side=tk.RIGHT, fill=tk.Y)

        self.layout_btn = tk.Button(
            right_frame, text="基础设置", bg="#555", fg="white",
            font=("微软雅黑", 11), padx=12, pady=4, relief=tk.RAISED, bd=2, command=self.open_layout_window
        )
        self.layout_btn.pack(side=tk.RIGHT, padx=3)

    def create_seat_container(self):
        # 创建外层容器，用于放置滚动条和座位框架
        self.seat_outer_container = tk.Frame(self.root, bg="#f8f8f8")
        self.seat_outer_container.pack(fill=tk.BOTH, expand=True, padx=8, pady=4)

        # 创建水平滚动条
        h_scrollbar = tk.Scrollbar(self.seat_outer_container, orient=tk.HORIZONTAL)
        h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)

        # 创建垂直滚动条
        v_scrollbar = tk.Scrollbar(self.seat_outer_container, orient=tk.VERTICAL)
        v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # 创建画布，用于放置座位框架并支持滚动
        self.seat_canvas = tk.Canvas(
            self.seat_outer_container, 
            bg="#f8f8f8",
            xscrollcommand=h_scrollbar.set,
            yscrollcommand=v_scrollbar.set,
            highlightthickness=0
        )
        self.seat_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # 配置滚动条与画布的关联
        h_scrollbar.config(command=self.seat_canvas.xview)
        v_scrollbar.config(command=self.seat_canvas.yview)

        # 创建座位框架，作为画布的子组件
        # 使用普通的tk.Frame替代ttk.LabelFrame，避免样式问题
        self.seat_frame = tk.LabelFrame(
            self.seat_canvas, text="教室座位布局",
            relief="solid",
            borderwidth=2,
            bg="#f8f8f8"
        )
        
        # 将座位框架添加到画布中，并使其居中显示
        self.seat_canvas.create_window((0, 0), window=self.seat_frame, anchor="nw")
        
        # 绑定大小变化事件，确保画布能够正确调整滚动区域
        self.seat_frame.bind("<Configure>", self.on_seat_frame_configure)
        
        # 绑定画布大小变化事件，确保座位框架始终居中
        self.seat_canvas.bind("<Configure>", self.on_canvas_configure)
        
        # 使画布可以通过鼠标滚轮滚动
        self.seat_canvas.bind_all("<MouseWheel>", self._on_mousewheel)
        
    def on_seat_frame_configure(self, event):
        """调整画布的滚动区域以适应座位框架的大小"""
        self.seat_canvas.configure(scrollregion=self.seat_canvas.bbox("all"))
    
    def on_canvas_configure(self, event):
        """确保座位框架在画布中居中显示"""
        # 获取画布和座位框架的尺寸
        canvas_width = event.width
        seat_frame_width = self.seat_frame.winfo_width()
        
        # 计算水平居中的位置
        x_pos = max(0, (canvas_width - seat_frame_width) // 2)
        
        # 更新座位框架在画布中的位置
        self.seat_canvas.coords(1, x_pos, 0)  # 1是create_window返回的ID
        
    def _on_mousewheel(self, event):
        """处理鼠标滚轮事件，实现垂直滚动"""
        self.seat_canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        # 阻止事件传播，避免影响其他控件
        return "break"
    
    def lighten_color(self, color, percent):
        """将颜色调亮指定百分比"""
        color = color.lstrip('#')
        r = int(color[0:2], 16)
        g = int(color[2:4], 16)
        b = int(color[4:6], 16)
        
        # 计算新的RGB值
        r = min(255, r + int((255 - r) * percent / 100))
        g = min(255, g + int((255 - g) * percent / 100))
        b = min(255, b + int((255 - b) * percent / 100))
        
        # 转换回十六进制格式
        return f'#{r:02x}{g:02x}{b:02x}'

    # ---------------------- 布局设置（讲台永远居中） ----------------------
    def open_layout_window(self):
        layout_win = tk.Toplevel(self.root)
        layout_win.title("基础设置")
        layout_win.geometry("350x300")
        layout_win.resizable(False, False)

        # 班级信息部分
        tk.Label(layout_win, text="班级名称：", font=("微软雅黑", 10)).grid(row=0, column=0, padx=10, pady=10, sticky="w")
        self.class_entry = ttk.Entry(layout_win, width=20)
        self.class_entry.grid(row=0, column=1, padx=10, pady=10, sticky="w")
        self.class_entry.insert(0, self.layout_config["class_name"])
        
        tk.Label(layout_win, text="班主任姓名：", font=("微软雅黑", 10)).grid(row=1, column=0, padx=10, pady=10, sticky="w")
        self.teacher_entry = ttk.Entry(layout_win, width=20)
        self.teacher_entry.grid(row=1, column=1, padx=10, pady=10, sticky="w")
        self.teacher_entry.insert(0, self.layout_config["teacher_name"])
        
        # 布局设置部分
        tk.Label(layout_win, text="讲台侧座位数：", font=("微软雅黑", 10)).grid(row=2, column=0, padx=10, pady=10, sticky="w")
        self.podium_entry = ttk.Entry(layout_win, width=10)
        self.podium_entry.grid(row=2, column=1, padx=10, pady=10)
        self.podium_entry.insert(0, str(self.layout_config["podium_seats"]))

        tk.Label(layout_win, text="主体座位行数：", font=("微软雅黑", 10)).grid(row=3, column=0, padx=10, pady=10, sticky="w")
        self.row_entry = ttk.Entry(layout_win, width=10)
        self.row_entry.grid(row=3, column=1, padx=10, pady=10)
        self.row_entry.insert(0, str(self.layout_config["main_rows"]))

        tk.Label(layout_win, text="主体座位列数：", font=("微软雅黑", 10)).grid(row=4, column=0, padx=10, pady=10, sticky="w")
        self.col_entry = ttk.Entry(layout_win, width=10)
        self.col_entry.grid(row=4, column=1, padx=10, pady=10)
        self.col_entry.insert(0, str(self.layout_config["main_cols"]))

        # 确认按钮
        ttk.Button(
            layout_win, text="确认设置",
            command=lambda: self.apply_layout(layout_win)
        ).grid(row=5, column=0, columnspan=2, pady=10)

    def save_config(self):
        """保存配置到config.ini文件"""
        config_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.ini")
        config = configparser.ConfigParser()
        config['Layout'] = {
            'class_name': self.layout_config['class_name'],
            'teacher_name': self.layout_config['teacher_name'],
            'podium_seats': str(self.layout_config['podium_seats']),
            'main_rows': str(self.layout_config['main_rows']),
            'main_cols': str(self.layout_config['main_cols'])
        }
        with open(config_file, 'w', encoding='utf-8') as f:
            config.write(f)
        # 配置更改后也自动保存座位数据
        self.auto_save_data()
    
    def apply_layout(self, layout_win):
        try:
            podium_seats = int(self.podium_entry.get())
            main_rows = int(self.row_entry.get())
            main_cols = int(self.col_entry.get())
            if podium_seats < 0 or podium_seats > 4 or main_rows < 1 or main_cols < 1:
                messagebox.showerror("错误", "行数/列数需≥1，讲台侧座位数需≥0且≤4")
                return
        except ValueError:
            messagebox.showerror("错误", "请输入数字")
            return

        # 获取班级和班主任信息
        class_name = self.class_entry.get().strip()
        teacher_name = self.teacher_entry.get().strip()

        self.layout_config = {
            "podium_seats": podium_seats,
            "main_rows": main_rows,
            "main_cols": main_cols,
            "class_name": class_name,
            "teacher_name": teacher_name
        }
        
        # 保存配置到config.ini
        self.save_config()
        
        self.generate_seat_positions()
        layout_win.destroy()
        self.seat_frame.config(
            text=f"教室座位布局（讲台侧{podium_seats}个座位 + {main_rows}×{main_cols}座位）"
        )
    def generate_seat_positions(self):
        self.seat_positions = []
        main_cols = self.layout_config["main_cols"]
        podium_seats = self.layout_config["podium_seats"]
        main_rows = self.layout_config["main_rows"]
        # 计算总宽度与讲台参数 - 讲台独立居中，不与座位列对齐
        total_width = max(main_cols, 4)  # 增加最小宽度以适应2格讲台
        podium_width = 2  # 讲台固定占用2格
        podium_start_col = (total_width - podium_width + 1) // 2  # 讲台始终居中，不考虑座位列
        podium_row = main_rows + 1  # 讲台行位置（教师视角：讲台在下方）

        # 1. 讲台侧座位（若数量>0则生成）- 围绕讲台左右排列
        # 教师视角：讲台侧座位位于讲台同一行的左右两侧
        if podium_seats > 0:
            # 计算左侧和右侧座位数
            left_seats = podium_seats // 2
            right_seats = podium_seats - left_seats
            
            # 左侧座位（从教师视角看是左侧，位于讲台左边）
            for i in range(left_seats):
                self.seat_positions.append((podium_row, podium_start_col - (left_seats - i)))
            
            # 右侧座位（从教师视角看是右侧，位于讲台右边）
            for i in range(right_seats):
                self.seat_positions.append((podium_row, podium_start_col + podium_width + i))

        # 2. 主体座位 - 教师视角：左下角为起始号
        # 从下往上遍历行，从左到右遍历列，确保左下角座位先被添加（起始号）
        start_col = (total_width - main_cols) // 2
        for r in range(self.layout_config["main_rows"], 0, -1):
            # 从左到右遍历列（教师视角）
            for c in range(start_col, start_col + main_cols):
                self.seat_positions.append((r, c))

        # 座位编号映射
        self.seat_index_map = {pos: idx + 1 for idx, pos in enumerate(self.seat_positions)}
        self.seat_data = {pos: {"name": "空", "gender": "空"} for pos in self.seat_positions}
        self.refresh_seat_buttons()

    def refresh_seat_buttons(self):
        # 清空原有控件
        for widget in self.seat_frame.winfo_children():
            widget.destroy()

        main_cols = self.layout_config["main_cols"]
        podium_seats = self.layout_config["podium_seats"]
        main_rows = self.layout_config["main_rows"]
        # 保持与generate_seat_positions方法中相同的变量名
        total_width = max(main_cols, 3)

        # 强制保留讲台并始终居中显示 - 教师视角：讲台位于教室前方（下方）
        total_width = max(main_cols, 4)  # 增加最小宽度以适应2格讲台
        podium_width = 2  # 讲台固定占用2格
        podium_start_col = (total_width - podium_width + 1) // 2  # 讲台始终居中，不考虑座位列
        podium_row = main_rows + 1  # 讲台行位置
        
        # 改进的讲台样式
        podium_frame = tk.Frame(
            self.seat_frame, 
            bg="#8E24AA",  # 更深的紫色，更显专业
            height=55,  # 增加高度
            bd=2,  # 添加边框
            relief=tk.RAISED  # 凸起效果
        )
        podium_frame.grid(
            row=podium_row, column=podium_start_col, columnspan=podium_width,
            pady=12, padx=8, sticky="nsew"  # 增加间距
        )
        
        # 改进的讲台标签
        podium_label = tk.Label(
            podium_frame, 
            text="讲 台", 
            font=("微软雅黑", 18, "bold"),  # 更大的字体
            fg="white", 
            bg="#8E24AA",
            bd=0  # 无边框
        )
        podium_label.pack(expand=True, fill=tk.BOTH, pady=4, padx=8)
        
        # 添加班级和教师信息显示
        if self.layout_config["class_name"] or self.layout_config["teacher_name"]:
            info_text = ""
            if self.layout_config["class_name"]:
                info_text += f"班级：{self.layout_config['class_name']}  "
            if self.layout_config["teacher_name"]:
                info_text += f"班主任：{self.layout_config['teacher_name']}"
            
            info_frame = tk.Frame(self.seat_frame, bg="#f8f8f8", height=30)
            info_frame.grid(
                row=0, column=0, columnspan=total_width,
                pady=5, padx=10, sticky="ew"
            )
            
            tk.Label(
                info_frame, 
                text=info_text, 
                font=("微软雅黑", 12, "italic"), 
                fg="#555555", 
                bg="#f8f8f8"
            ).pack(side=tk.TOP, expand=True)  # 使用TOP而不是CENTER

        # 生成座位按钮
        self.seat_buttons = {}
        for pos in self.seat_positions:
            r, c = pos
            seat_idx = self.seat_index_map[pos]
            student_info = self.seat_data[pos]
            
            # 创建座位按钮 - 原始版本
            btn = tk.Button(
                self.seat_frame, 
                text=f"{seat_idx}\n{student_info['name']}",
                width=7,  # 原始宽度设置
                height=3,  # 原始高度设置
                font=('微软雅黑', 9, 'bold'),  # 原始字体大小
                bg=self.gender_color[student_info['gender']],  # 原始颜色设置
                fg="#333333",
                relief=tk.FLAT,  # 原始边框样式
                bd=1,
                cursor="hand2"
            )
            
            # 使用原始的grid布局设置
            btn.grid(row=r, column=c, padx=5, pady=3)
            
            # 存储位置信息
            btn.pos_info = pos
            self.seat_buttons[pos] = btn
            
            # 绑定拖拽事件
            btn.bind("<Button-1>", lambda e, p=pos: self.on_drag_start(e, p))
            btn.bind("<B1-Motion>", lambda e, p=pos: self.on_drag_motion(e, p))
            btn.bind("<ButtonRelease-1>", lambda e, p=pos: self.on_drag_end(e, p))
            
            # 鼠标悬停效果 - 恢复原始功能
            btn.bind("<Enter>", lambda e: e.widget.config(bg=self.lighten_color(e.widget.cget("bg"), 0.1)))
            btn.bind("<Leave>", lambda e: e.widget.config(bg=self.gender_color[student_info['gender']]))

    # ---------------------- 原有功能保留 ----------------------
    def import_excel(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel文件", "*.xlsx;*.xls")])
        if not file_path:
            return
        try:
            df = pd.read_excel(file_path)
            required_cols = ["姓名", "性别", "身高"]
            if not all(col in df.columns for col in required_cols):
                messagebox.showerror("错误", "Excel需包含：姓名、性别、身高列")
                return
            self.students = df.to_dict("records")
            messagebox.showinfo("成功", f"已导入{len(self.students)}名学生")
            # 导入数据后自动保存
            self.auto_save_data()
        except Exception as e:
            messagebox.showerror("导入失败", str(e))

    def random_arrange(self):
        if not self.students:
            messagebox.showwarning("提示", "请先导入学生数据")
            return
        self.seat_data = {pos: {"name": "空", "gender": "空"} for pos in self.seat_positions}
        random_students = random.sample(self.students, min(len(self.students), len(self.seat_positions)))
        for pos, stu in zip(self.seat_positions, random_students):
            self.seat_data[pos] = {"name": stu["姓名"], "gender": stu["性别"]}
        self.update_seat_buttons()
        # 随机排列后自动保存
        self.auto_save_data()

    def sort_by_height(self):
        if not self.students:
            messagebox.showwarning("提示", "请先导入学生数据")
            return
        sorted_stu = sorted(self.students, key=lambda x: x["身高"])
        self.seat_data = {pos: {"name": "空", "gender": "空"} for pos in self.seat_positions}
        for pos, stu in zip(self.seat_positions, sorted_stu):
            self.seat_data[pos] = {"name": stu["姓名"], "gender": stu["性别"]}
        self.update_seat_buttons()
        # 排序后自动保存
        self.auto_save_data()
        
    def sort_by_score(self):
        if not self.students:
            messagebox.showwarning("提示", "请先导入学生数据")
            return
        
        # 检查是否有学生包含成绩字段
        has_score = any("成绩" in student for student in self.students)
        if not has_score:
            messagebox.showwarning("提示", "当前学生数据中没有成绩信息")
            return
        
        # 询问用户排序方式
        from tkinter import simpledialog
        result = simpledialog.askstring("排序方式", "请选择排序方式：\n1. 成绩从高到低\n2. 成绩从低到高", 
                                      initialvalue="1")
        
        if result not in ["1", "2"]:
            return
        
        # 根据选择进行排序
        reverse = result == "1"  # 1表示从高到低，2表示从低到高
        
        # 排序学生，使用默认值处理没有成绩的学生
        sorted_stu = sorted(
            self.students,
            key=lambda x: x.get("成绩", 0),
            reverse=reverse
        )
        
        # 清空座位数据并重新分配
        self.seat_data = {pos: {"name": "空", "gender": "空"} for pos in self.seat_positions}
        for pos, stu in zip(self.seat_positions, sorted_stu):
            self.seat_data[pos] = {"name": stu["姓名"], "gender": stu["性别"]}
        
        # 更新界面和保存数据
        self.update_seat_buttons()
        self.auto_save_data()
        
        # 显示排序成功的提示
        order_text = "从高到低" if reverse else "从低到高"
        messagebox.showinfo("成功", f"已按成绩{order_text}排序")

    def on_drag_start(self, event, pos):
        # 记录拖拽源信息
        self.drag_source = pos
        # 更改拖拽按钮的外观
        self.seat_buttons[pos].config(relief=tk.SUNKEN, bg="#FFA726")
        # 记录鼠标按下的位置（相对于按钮）
        self.drag_x_offset = event.x
        self.drag_y_offset = event.y
    
    def on_drag_motion(self, event, pos):
        # 只有当当前按钮是拖拽源时才处理移动
        if self.drag_source == pos:
            btn = self.seat_buttons[pos]
            # 先将按钮从grid布局中移除，改用place布局
            btn.grid_forget()
            # 获取窗口中的绝对位置
            x = event.x_root - self.drag_x_offset - self.seat_frame.winfo_rootx()
            y = event.y_root - self.drag_y_offset - self.seat_frame.winfo_rooty()
            # 更新按钮位置（相对于seat_frame）
            btn.place(x=x, y=y)
    
    def on_drag_end(self, event, pos):
        # 确保拖拽源存在
        if self.drag_source:
            # 获取鼠标释放时的实际窗口坐标
            x, y = self.seat_frame.winfo_pointerxy()
            # 转换为seat_frame内部坐标
            x_rel = x - self.seat_frame.winfo_rootx()
            y_rel = y - self.seat_frame.winfo_rooty()
            
            # 查找鼠标释放位置下的目标按钮
            target_pos = None
            for btn_pos, btn in self.seat_buttons.items():
                # 获取按钮的边界框
                btn_x1 = btn.winfo_x()
                btn_y1 = btn.winfo_y()
                btn_x2 = btn_x1 + btn.winfo_width()
                btn_y2 = btn_y1 + btn.winfo_height()
                
                # 检查鼠标是否在按钮范围内
                if btn_x1 <= x_rel <= btn_x2 and btn_y1 <= y_rel <= btn_y2:
                    target_pos = btn_pos
                    break
            
            # 如果找到目标按钮且不是拖拽源本身，则交换数据
            if target_pos and target_pos != self.drag_source:
                # 交换座位数据
                temp = self.seat_data[self.drag_source]
                self.seat_data[self.drag_source] = self.seat_data[target_pos]
                self.seat_data[target_pos] = temp
            
            # 无论是否交换，都需要重置所有按钮布局
            self.update_seat_buttons()
            # 重置拖拽源
            self.drag_source = None
            # 拖拽结束后自动保存
            self.auto_save_data()

    def update_seat_buttons(self):
        for pos in self.seat_positions:
            seat_idx = self.seat_index_map[pos]
            data = self.seat_data[pos]
            btn = self.seat_buttons[pos]
            # 重置按钮状态和样式
            btn.config(
                text=f"{seat_idx}\n{data['name']}",  # 移除性别信息
                bg=self.gender_color[data["gender"]],
                relief=tk.RAISED
            )
            # 确保按钮回到grid布局
            r, c = pos
            # 首先取消place布局，然后重新应用grid布局
            btn.place_forget()
            btn.grid(row=r, column=c, padx=10, pady=8)

    def show_seat_info(self, pos):
        data = self.seat_data[pos]
        if data["name"] == "空":
            messagebox.showinfo("座位信息", f"座位{self.seat_index_map[pos]}：空")
        else:
            height = next((s["身高"] for s in self.students if s["姓名"] == data["name"]), "未知")
            messagebox.showinfo(
                "座位信息",
                f"座位{self.seat_index_map[pos]}\n姓名：{data['name']}\n性别：{data['gender']}\n身高：{height}cm"
            )

    def reset_seats(self):
        if messagebox.askyesno("确认", "确定重置所有座位吗？"):
            self.seat_data = {pos: {"name": "空", "gender": "空"} for pos in self.seat_positions}
            self.update_seat_buttons()
            # 重置座位后自动保存
            self.auto_save_data()
    
    def auto_save_data(self):
        """自动保存座位表数据到根目录的JSON文件，不弹出对话框"""
        file_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "座位表数据.json")
        
        try:
            # 准备要保存的数据
            # 创建学生列表
            students_without_scores = [{key: value for key, value in student.items() if key != "成绩"} for student in self.students]
            
            save_data = {
                "layout_config": self.layout_config,
                "seat_data": {str(pos): data for pos, data in self.seat_data.items()},
                "seat_index_map": {str(pos): idx for pos, idx in self.seat_index_map.items()},
                "students": students_without_scores
            }
            
            # 获取程序运行时的基础目录
            if hasattr(sys, '_MEIPASS'):
                base_dir = os.path.dirname(os.path.abspath(sys.executable))
            else:
                base_dir = os.path.dirname(os.path.abspath(__file__))
                
            file_path = os.path.join(base_dir, "座位表数据.json")
            # 保存到文件
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(save_data, f, ensure_ascii=False, indent=2)
            
            # 自动保存不显示提示，避免干扰用户
            # print(f"数据已自动保存到：{file_path}")  # 调试时可以启用
        except IOError as e:
            # 处理文件IO错误
            print(f"自动保存失败（IO错误）：{str(e)}")
        except json.JSONDecodeError as e:
            # 处理JSON解析错误
            print(f"自动保存失败（JSON解析错误）：{str(e)}")
        except Exception as e:
            # 处理其他未知错误
            print(f"自动保存失败（未知错误）：{str(e)}")
    
    def auto_load_data(self):
        """程序启动时自动尝试读取座位表数据文件"""
        # 获取程序运行时的基础目录
        if hasattr(sys, '_MEIPASS'):
            base_dir = os.path.dirname(os.path.abspath(sys.executable))
        else:
            base_dir = os.path.dirname(os.path.abspath(__file__))
            
        file_path = os.path.join(base_dir, "座位表数据.json")
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            # 文件不存在时，创建初始空座位数据
            self.seat_data = {pos: {"name": "空", "gender": "空"} for pos in self.seat_positions}
            return
        
        try:
            # 读取文件
            with open(file_path, 'r', encoding='utf-8') as f:
                load_data = json.load(f)
            
            # 验证必要字段
            required_fields = ["layout_config", "seat_data", "seat_index_map", "students"]
            if not all(field in load_data for field in required_fields):
                # 文件格式无效时，创建初始空座位数据
                self.seat_data = {pos: {"name": "空", "gender": "空"} for pos in self.seat_positions}
                return
            
            # 恢复数据
            # 使用try-except避免数据类型不匹配导致的程序崩溃
            try:
                # 如果加载的布局配置与当前配置不同，保留当前配置
                # 只更新座位数据、学生信息等
                self.students = load_data["students"]
                
                # 使用类方法安全地解析元组字符串
                loaded_seat_data = {}
                loaded_seat_index_map = {}
                for pos_str, data in load_data["seat_data"].items():
                    pos = self.parse_tuple_str(pos_str)
                    if pos:
                        loaded_seat_data[pos] = data
                
                for pos_str, idx in load_data["seat_index_map"].items():
                    pos = self.parse_tuple_str(pos_str)
                    if pos:
                        loaded_seat_index_map[pos] = idx
                
                # 只保留当前座位布局中存在的位置数据
                self.seat_data = {}
                for pos in self.seat_positions:
                    if pos in loaded_seat_data:
                        self.seat_data[pos] = loaded_seat_data[pos]
                    else:
                        self.seat_data[pos] = {"name": "空", "gender": "空"}
                
                # 仅在索引映射有效时更新
                if all(pos in loaded_seat_index_map for pos in self.seat_positions):
                    self.seat_index_map = loaded_seat_index_map
                
                # 更新座位按钮显示
                self.update_seat_buttons()
                
                # 静默加载，不显示提示
                # print(f"已自动加载上次保存的数据，共{len(self.students)}名学生")  # 调试时可以启用
            except (ValueError, IndexError, TypeError) as e:
                # 处理数据类型不匹配、索引错误或类型错误
                print(f"数据恢复失败：{str(e)}")
                self.seat_data = {pos: {"name": "空", "gender": "空"} for pos in self.seat_positions}
                
        except IOError as e:
            # 处理文件IO错误
            print(f"自动加载数据失败（IO错误）：{str(e)}")
            self.seat_data = {pos: {"name": "空", "gender": "空"} for pos in self.seat_positions}
        except json.JSONDecodeError as e:
            # 处理JSON解析错误
            print(f"自动加载数据失败（JSON解析错误）：{str(e)}")
            self.seat_data = {pos: {"name": "空", "gender": "空"} for pos in self.seat_positions}
        except Exception as e:
            # 处理其他未知错误
            print(f"自动加载数据失败（未知错误）：{str(e)}")
            self.seat_data = {pos: {"name": "空", "gender": "空"} for pos in self.seat_positions}
    
    def capture_seat_layout(self):
        """将座位表布局转换为图片
        
        Returns:
            PIL.Image: 座位表布局的图像对象，如果失败则返回None
        """
        try:
            # 检查PIL库是否可用
            if not pillow_available:
                print("PIL库不可用")
                return None
            
            # 获取座位区域的大小和布局信息
            # 创建一个适当大小的图像（基于座位数量和布局）
            rows = self.layout_config['main_rows']
            cols = self.layout_config['main_cols']
            
            # 计算图像尺寸（每个座位100x100像素，留出边距）
            # 根据座位数量动态调整座位大小，确保图像不会过大
            base_seat_size = 100
            max_seats_per_row = 10
            
            # 如果列数或行数较多，缩小座位尺寸
            if cols > max_seats_per_row:
                seat_size = int(base_seat_size * max_seats_per_row / cols)
            else:
                seat_size = base_seat_size
                
            # 限制最大尺寸，防止内存问题
            max_size = 3000
            margin = 150  # 留出顶部空间显示标题和底部空间显示讲台
            
            # 计算初始尺寸
            img_width = cols * seat_size + margin * 2
            img_height = rows * seat_size + margin * 2
            
            # 如果图像尺寸过大，等比例缩小
            scale_factor = 1.0
            if img_width > max_size or img_height > max_size:
                scale_factor = max_size / max(img_width, img_height)
                seat_size = int(seat_size * scale_factor)
                img_width = int(img_width * scale_factor)
                img_height = int(img_height * scale_factor)
            
            # 创建白色背景图像
            image = Image.new('RGB', (img_width, img_height), color='white')
            draw = ImageDraw.Draw(image)
            
            # 尝试加载中文字体（尝试多个常见的中文字体）
            font_paths = ["simhei.ttf", "simkai.ttf", "simsun.ttc", "msyh.ttc", "Arial.ttf"]
            
            def load_font(size, default=False):
                if default:
                    return ImageFont.load_default()
                
                for font_path in font_paths:
                    try:
                        return ImageFont.truetype(font_path, size)
                    except (IOError, OSError):
                        continue
                
                # 如果所有字体都加载失败，使用默认字体
                return ImageFont.load_default()
            
            # 加载不同大小的字体
            font = load_font(16)
            title_font = load_font(24)
            small_font = load_font(12)
            
            # 添加标题
            config = configparser.ConfigParser()
            config_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.ini")
            class_name = ""  # 默认值
            head_teacher = ""  # 默认值
            
            if os.path.exists(config_file):
                config.read(config_file, encoding='utf-8')
                if '班级信息' in config:
                    class_name = config['班级信息'].get('班级名称', '')
                    head_teacher = config['班级信息'].get('班主任', '')
            
            # 添加班级信息标题
            title_text = f"{class_name}"
            if head_teacher:
                title_text += f" - 班主任：{head_teacher}"
            
            # 计算文本尺寸并居中
            if hasattr(font, 'getsize'):
                title_width = font.getsize(title_text)[0]
            else:
                title_width = draw.textlength(title_text, font=title_font)
            
            draw.text((img_width // 2 - title_width // 2, margin // 3), title_text, font=title_font, fill='black')
            
            # 绘制座位
            for pos, data in self.seat_data.items():
                row, col = pos
                # 计算座位在图像中的位置
                x = margin + col * seat_size
                y = margin + row * seat_size
                
                # 绘制座位背景（根据性别设置不同颜色）
                if data["gender"] == "男":
                    fill_color = (220, 240, 255)  # 浅蓝色
                elif data["gender"] == "女":
                    fill_color = (255, 220, 230)  # 浅粉色
                else:
                    fill_color = (240, 240, 240)  # 浅灰色
                
                draw.rectangle([x, y, x + seat_size - 5, y + seat_size - 5], fill=fill_color, outline='black')
                
                # 添加座位号
                seat_idx = self.seat_index_map.get(pos, "")
                if seat_idx:
                    if hasattr(font, 'getsize'):
                        idx_width = font.getsize(str(seat_idx))[0]
                    else:
                        idx_width = draw.textlength(str(seat_idx), font=small_font)
                    draw.text((x + 5, y + 5), str(seat_idx), font=small_font, fill='black')
                
                # 添加学生姓名
                name = data["name"]
                if name != "空":
                    if hasattr(font, 'getsize'):
                        name_width = font.getsize(name)[0]
                    else:
                        name_width = draw.textlength(name, font=font)
                    draw.text((x + seat_size // 2 - name_width // 2, y + seat_size // 2 - 8), name, font=font, fill='black')
            
            # 绘制讲台（位于底部）
            podium_height = 40
            podium_y = img_height - margin + 20
            draw.rectangle([img_width // 4, podium_y, img_width * 3 // 4, podium_y + podium_height], fill='lightgray', outline='black')
            draw.text((img_width // 2 - 20, podium_y + 10), "讲台", font=font, fill='black')
            
            return image
        except Exception as e:
            print(f"创建座位布局图片失败：{str(e)}")
            return None
    
    def save_data(self):
        """保存座位布局和学生信息到本地JSON文件（用户手动保存）"""
        file_path = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("JSON文件", "*.json")],
            initialfile="座位表数据.json"
        )
        if not file_path:
            return
        
        try:
            # 准备要保存的数据
            # 创建学生列表
            students_without_scores = [{key: value for key, value in student.items() if key != "成绩"} for student in self.students]
            
            save_data = {
                "layout_config": self.layout_config,
                "seat_data": {str(pos): data for pos, data in self.seat_data.items()},
                "seat_index_map": {str(pos): idx for pos, idx in self.seat_index_map.items()},
                "students": students_without_scores
            }
            
            # 保存到文件
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(save_data, f, ensure_ascii=False, indent=2)
            
            messagebox.showinfo("成功", "数据已成功保存")
        except IOError as e:
            # 处理文件IO错误
            messagebox.showerror("保存错误", f"文件操作失败：{str(e)}")
        except json.JSONDecodeError as e:
            # 处理JSON解析错误
            messagebox.showerror("保存错误", f"JSON格式化失败：{str(e)}")
        except Exception as e:
            # 处理其他未知错误
            messagebox.showerror("保存错误", f"未知错误：{str(e)}")
    
    def load_data(self):
        """从本地JSON文件加载座位布局和学生信息"""
        file_path = filedialog.askopenfilename(
            filetypes=[("JSON文件", "*.json")]
        )
        if not file_path:
            return
        
        try:
            # 读取文件
            with open(file_path, 'r', encoding='utf-8') as f:
                load_data = json.load(f)
            
            # 验证必要字段
            required_fields = ["layout_config", "seat_data", "seat_index_map", "students"]
            if not all(field in load_data for field in required_fields):
                messagebox.showerror("错误", "无效的数据文件格式")
                return
            
            # 恢复数据
            self.layout_config = load_data["layout_config"]
            # 使用类方法安全地解析元组字符串
            self.seat_data = {}
            self.seat_index_map = {}
            for pos_str, data in load_data["seat_data"].items():
                pos = self.parse_tuple_str(pos_str)
                if pos:
                    self.seat_data[pos] = data
            
            for pos_str, idx in load_data["seat_index_map"].items():
                pos = self.parse_tuple_str(pos_str)
                if pos:
                    self.seat_index_map[pos] = idx
            self.students = load_data["students"]
            
            # 更新座位布局
            self.generate_seat_positions()
            # 重新应用座位数据
            # 使用类方法安全地解析元组字符串
            self.seat_data = {}
            self.seat_index_map = {}
            for pos_str, data in load_data["seat_data"].items():
                pos = self.parse_tuple_str(pos_str)
                if pos:
                    self.seat_data[pos] = data
            
            for pos_str, idx in load_data["seat_index_map"].items():
                pos = self.parse_tuple_str(pos_str)
                if pos:
                    self.seat_index_map[pos] = idx
            self.update_seat_buttons()
            
            messagebox.showinfo("成功", "数据已成功加载")
            # 加载数据后自动保存（确保数据一致性）
            self.auto_save_data()
        except IOError as e:
            # 处理文件IO错误
            messagebox.showerror("加载错误", f"文件操作失败：{str(e)}")
        except json.JSONDecodeError as e:
            # 处理JSON解析错误
            messagebox.showerror("加载错误", f"JSON解析失败：{str(e)}")
        except ValueError as e:
            # 处理数据格式错误
            messagebox.showerror("加载错误", f"数据格式错误：{str(e)}")
        except Exception as e:
            # 处理其他未知错误
            messagebox.showerror("加载错误", f"未知错误：{str(e)}")

    def export_pdf(self):
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import A4
            from reportlab.lib import colors
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
            from reportlab.lib.units import cm
            # 添加字体配置
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            # 添加样式导入
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        except ImportError:
            messagebox.showerror("错误", "请先安装reportlab：pip install reportlab")
            return
        
        # 从布局配置中获取班级和班主任信息（非必填）
        class_info = self.layout_config.get("class_name", "")
        teacher_info = self.layout_config.get("teacher_name", "")
        
        file_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF", "*.pdf")])
        if not file_path:
            return
        
        # 配置中文字体
        try:
            # 尝试注册Windows系统中的宋体
            pdfmetrics.registerFont(TTFont('SimSun', 'simsun.ttc'))
            # 尝试注册Windows系统中的微软雅黑
            pdfmetrics.registerFont(TTFont('MicrosoftYaHei', 'msyh.ttc'))
        except:
            # 如果注册失败，使用reportlab默认字体
            pass
        
        # 创建PDF文档 - 使用纵向A4页面
        doc = SimpleDocTemplate(file_path, pagesize=A4)
        elements = []
        
        # 添加标题
        styles = getSampleStyleSheet()
        # 设置标题样式：微软雅黑，小初大小（约36pt），居中对齐
        try:
            styles['Title'].fontName = 'MicrosoftYaHei'
        except:
            # 如果微软雅黑不可用，使用宋体作为备选
            styles['Title'].fontName = 'SimSun'
        
        styles['Title'].fontSize = 36  # 小初大小约为36pt
        styles['Title'].alignment = 1  # 1表示居中对齐
        
        # 设置正文样式为宋体四号（14pt）
        try:
            styles['BodyText'].fontName = 'SimSun'
        except:
            pass  # 如果字体不可用，使用默认字体
        styles['BodyText'].fontSize = 14  # 四号字体约为14pt
        styles['BodyText'].alignment = 0  # 0表示居左对齐（修复备注居左显示）
        
        # 添加座位表标题（居中显示，每个字间隔指定空格数，从配置文件读取设置）
        main_title = self.config.get("Export", "main_title", fallback="座位表")
        space_count = self.config.getint("Export", "title_space_count", fallback=2)
        
        spaced_title = " ".join(main_title)  # 为每个字添加一个空格间隔
        # 根据配置的空格数添加额外空格
        if space_count > 1:
            spaced_title = spaced_title.replace(" ", " " * space_count)
        
        # 设置标题为指定字体、大小、加粗、黑色文本
        title_text_color = self.config.get("Color", "title_text_color", fallback="black")
        if hasattr(colors, title_text_color):
            styles['Title'].textColor = getattr(colors, title_text_color)
        else:
            styles['Title'].textColor = colors.black  # 默认黑色
            
        # 设置标题下划线颜色
        title_underline_color = self.config.get("Color", "title_underline_color", fallback="black")
        if hasattr(colors, title_underline_color):
            styles['Title'].textUnderlineColor = getattr(colors, title_underline_color)
        else:
            styles['Title'].textUnderlineColor = colors.black  # 默认黑色
            
        title = Paragraph(f"<b><u>{spaced_title}</u></b>", styles['Title'])
        elements.append(title)
        
        # 添加空行作为标题和班级信息之间的间距，避免标题遮挡班级信息
        elements.append(Spacer(1, 0.3*cm))
        
        # 添加班级和班主任信息（仅在有信息时显示，居中显示，班级和班主任姓名带下划线）
        if class_info or teacher_info:
            info_parts = []
            if class_info:
                # 班级名称带下划线
                info_parts.append(f"班级：<u>{class_info}</u>")
            if teacher_info:
                # 班主任姓名带下划线
                info_parts.append(f"班主任：<u>{teacher_info}</u>")
            info_text = "  ".join(info_parts)
            # 创建一个居中对齐的样式来显示班级信息
            info_style = ParagraphStyle('InfoText', parent=styles['BodyText'])
            info_style.alignment = 1  # 1表示居中对齐
            info_paragraph = Paragraph(info_text, info_style)
            elements.append(info_paragraph)
        
        # 添加空行
        elements.append(Spacer(1, 0.5*cm))
        
        # 使用座位布局图片替代表格
        image_success = False
        
        try:
            # 生成座位布局图片
            seat_image = self.capture_seat_layout()
            
            if seat_image:
                try:
                    # 临时保存图片到内存
                    import io
                    img_buffer = io.BytesIO()
                    seat_image.save(img_buffer, format='PNG')
                    img_buffer.seek(0)
                    
                    # 计算图片大小，使其适应A4页面（减去边距）
                    page_width, page_height = A4  # A4尺寸
                    margin = 2 * cm  # 边距
                    available_width = page_width - 2 * margin  # 可用宽度
                    available_height = page_height - 6 * cm  # 可用高度（考虑标题、备注等内容）
                    
                    # 计算图片尺寸，保持原始宽高比
                    img_width, img_height = seat_image.size
                    aspect_ratio = img_height / img_width
                    
                    # 首先尝试按宽度缩放
                    display_width = available_width
                    display_height = display_width * aspect_ratio
                    
                    # 如果高度超出可用空间，则按高度缩放
                    if display_height > available_height:
                        display_height = available_height
                        display_width = display_height / aspect_ratio
                    
                    # 确保图片不会太小
                    min_size = 5 * cm
                    if display_width < min_size or display_height < min_size:
                        scale_factor = min_size / min(display_width, display_height)
                        display_width *= scale_factor
                        display_height *= scale_factor
                    
                    # 添加图片到PDF
                    rl_image = RLImage(img_buffer, width=display_width, height=display_height)
                    elements.append(rl_image)
                    image_success = True
                except Exception as e:
                    print(f"添加座位布局图片到PDF失败：{str(e)}")
                    # 如果图片添加失败，显示错误信息
                    error_text = Paragraph(f"<b>座位布局图片添加失败：{str(e)}</b>", styles['BodyText'])
                    elements.append(error_text)
            else:
                # 如果图片生成失败，显示错误信息
                error_text = Paragraph("<b>座位布局图片生成失败，请检查座位数据</b>", styles['BodyText'])
                elements.append(error_text)
                
            # 如果图片添加成功，添加空行
            if image_success:
                elements.append(Spacer(1, 0.5*cm))
                
        except Exception as e:
            # 捕获任何未预期的错误
            print(f"处理座位布局图片时发生错误：{str(e)}")
            error_text = Paragraph(f"<b>处理座位布局图片时发生错误：{str(e)}</b>", styles['BodyText'])
            elements.append(error_text)
        
        # 添加备注信息
        elements.append(Spacer(1, 1*cm))  # 添加垂直间距
        note_text = Paragraph("<b>备注：</b>", styles['BodyText'])
        elements.append(note_text)
        note_content1 = Paragraph("1、座位安排主要依据为身高，同时参考学生性别、性格、学习成绩等因素进行互补性编排；", styles['BodyText'])
        elements.append(note_content1)
        note_content2 = Paragraph("2、班级座位每月根据实际情况调整。", styles['BodyText'])
        elements.append(note_content2)
        
        # 构建PDF文档
        try:
            doc.build(elements)
            messagebox.showinfo("成功", "座位布局已导出为PDF")
        except Exception as e:
            messagebox.showerror("导出错误", f"PDF导出失败：{str(e)}")
    
    def export_layout_to_word(self):
        if not docx_available:
            messagebox.showerror("错误", "请先安装python-docx：pip install python-docx")
            return
        
        # 导入必要的docx模块
        from docx.shared import Inches, Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_ALIGN_VERTICAL
        
        # 从布局配置中获取班级和班主任信息（非必填）
        class_info = self.layout_config.get("class_name", "")
        teacher_info = self.layout_config.get("teacher_name", "")
        
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word文档", "*.docx")])
        if not file_path:
            return
        
        doc = Document()
        
        # 添加居中的座位表标题，设置为微软雅黑，小初大小（约36pt）
        title = doc.add_heading('', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 为标题文本添加两个空格间隔
        title_text = "座位表"
        spaced_title = "  ".join(title_text)  # 每个字之间添加两个空格
        
        # 创建标题run并设置字体样式：微软雅黑，小初大小，加粗，黑色文本，下划线为黑色填充
        title_run = title.add_run(spaced_title)
        title_run.font.name = '微软雅黑'
        title_run.font.size = Pt(36)  # 小初大小约为36pt
        title_run.bold = True
        # 设置文本颜色为黑色
        from docx.shared import RGBColor
        title_run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
        # 设置下划线为黑色填充
        title_run.underline = True
        # 确保下划线颜色为黑色（python-docx默认下划线颜色与文本颜色相同）
        
        # 添加空行作为标题和班级信息之间的间距，避免标题遮挡班级信息
        doc.add_paragraph()
        
        # 添加班级和班主任信息（仅在有信息时显示，居中显示，班级和班主任姓名带下划线）
        if class_info or teacher_info:
            info_paragraph = doc.add_paragraph()
            info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 确保班级信息居中显示
            
            # 分别处理班级和班主任信息，为姓名添加下划线，所有文本设置为宋体四号
            if class_info:
                run1 = info_paragraph.add_run("班级：")
                run1.font.name = '宋体'
                run1.font.size = Pt(14)  # 四号字体
                # 班级名称带下划线
                class_run = info_paragraph.add_run(class_info)
                class_run.underline = True
                class_run.font.name = '宋体'
                class_run.font.size = Pt(14)  # 四号字体
                
            if class_info and teacher_info:
                run_space = info_paragraph.add_run("  ")
                run_space.font.name = '宋体'
                run_space.font.size = Pt(14)  # 四号字体
            
            if teacher_info:
                run2 = info_paragraph.add_run("班主任：")
                run2.font.name = '宋体'
                run2.font.size = Pt(14)  # 四号字体
                # 班主任姓名带下划线
                teacher_run = info_paragraph.add_run(teacher_info)
                teacher_run.underline = True
                teacher_run.font.name = '宋体'
                teacher_run.font.size = Pt(14)  # 四号字体

        # 获取布局配置参数
        main_rows = self.layout_config["main_rows"]
        main_cols = self.layout_config["main_cols"]
        podium_seats = self.layout_config["podium_seats"]
        
        # 计算总宽度和讲台参数，与generate_seat_positions方法保持一致
        total_width = max(main_cols, 4)  # 增加最小宽度以适应2格讲台
        podium_width = 2  # 讲台固定占用2格
        podium_start_col = (total_width - podium_width + 1) // 2  # 讲台始终居中
        podium_row = main_rows + 1  # 讲台行位置（教师视角：讲台在下方）
        
        # 创建一个表格来表示座位布局，包含所有需要的行和列
        table = doc.add_table(rows=podium_row + 1, cols=total_width)
        table.style = 'Table Grid'  # 设置表格样式为网格
        
        # 设置表格为自动调整以适应窗口宽度
        # 使用try-except块确保兼容性
        try:
            # 获取表格的属性
            tbl = table._tbl
            
            # 使用更兼容的方式设置表格宽度属性
            if hasattr(tbl, 'get_or_add_tblPr'):
                tblPr = tbl.get_or_add_tblPr()
            else:
                # 尝试其他可能的方法或直接使用XML操作
                from docx.oxml.shared import OxmlElement, qn
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            # 添加自动调整属性
            if hasattr(tblPr, 'add_tblW'):
                tblW = tblPr.add_tblW()
            else:
                tblW = OxmlElement('w:tblW')
                tblPr.append(tblW)
            
            tblW.set(qn('w:type'), 'auto')
            tblW.set(qn('w:w'), '0')
        except Exception as e:
            # 记录错误但不中断程序执行
            print(f"设置表格自动调整属性时出错: {e}")

        # 填充讲台信息 - 独立居中显示
        podium_cell = table.cell(podium_row, podium_start_col)
        podium_cell.merge(table.cell(podium_row, podium_start_col + podium_width - 1))
        
        # 设置讲台单元格样式：宋体四号，行高1.5CM，垂直居中
        podium_cell.text = ""
        podium_run = podium_cell.paragraphs[0].add_run("讲台")
        podium_run.font.name = '宋体'
        podium_run.font.size = Pt(14)
        podium_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        podium_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        podium_cell.height = Cm(1.5)
        
        # 设置讲台单元格背景色以突出显示
        podium_run.font.color.rgb = RGBColor(255, 255, 255)  # 白色文字
        shading_elm = parse_xml(r'<w:shd {} w:fill="9C27B0"/>'.format(nsdecls('w')))
        podium_cell._tc.get_or_add_tcPr().append(shading_elm)

        # 设置所有单元格格式：宋体四号，垂直居中，行高1CM
        
        # 遍历所有单元格设置格式
        for row_idx in range(podium_row + 1):
            for col_idx in range(total_width):
                cell = table.cell(row_idx, col_idx)
                # 设置单元格垂直居中
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # 设置单元格行高为1.5CM
                cell.height = Cm(1.5)
                # 设置段落水平居中
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # 设置字体为宋体四号（约14pt）
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run.font.size = Pt(14)
                    # 如果段落没有run，添加一个默认run
                    if not cell.paragraphs[0].runs:
                        run = cell.paragraphs[0].add_run()
                        run.font.name = '宋体'
                        run.font.size = Pt(14)
        
        # 填充座位信息
        for pos in self.seat_positions:
            r, c = pos
            try:
                # 确保在表格范围内
                if 0 <= r <= podium_row and 0 <= c < total_width:
                    cell = table.cell(r, c)
                    data = self.seat_data[pos]
                    # 任务3：更新导出逻辑，使座位有学生时显示名字，无学生时显示为空
                    if data['name'] and data['name'] != "空":
                        # 清除现有内容，添加新的格式化文本
                        cell.text = ""
                        run = cell.paragraphs[0].add_run(data['name'])
                        run.font.name = '宋体'
                        run.font.size = Pt(14)
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        cell.text = ""  # 空白单元格
            except Exception:
                pass  # 忽略单元格填充错误
        
        # 添加备注信息
        # 添加空行作为间距
        doc.add_paragraph()
        
        # 添加备注标题（加粗，居左显示）
        note_title = doc.add_paragraph()
        note_run = note_title.add_run("备注：")
        note_run.bold = True
        note_title.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 设置备注标题居左
        
        # 添加备注内容1（居左显示）
        note_content1 = doc.add_paragraph("1、座位安排主要依据为身高，同时参考学生性别、性格、学习成绩等因素进行互补性编排；")
        note_content1.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 设置备注内容1居左
        
        # 添加备注内容2（居左显示）
        note_content2 = doc.add_paragraph("2、班级座位每月根据实际情况调整。")
        note_content2.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 设置备注内容2居左
        
        try:
            doc.save(file_path)
            messagebox.showinfo("成功", "座位布局已导出为Word文档")
        except Exception as e:
            messagebox.showerror("导出错误", f"Word文档导出失败：{str(e)}")

if __name__ == "__main__":
    root = tk.Tk()
    app = StudentSeatTool(root)
    root.mainloop()